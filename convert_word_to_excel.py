"""
Скрипт для конвертации Word документа в Excel формат
"""
import sys
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import re
from datetime import datetime


def extract_text_from_word(word_path: str) -> list:
    """
    Извлечение текста из Word документа
    
    Args:
        word_path: Путь к Word файлу
    
    Returns:
        Список строк текста
    """
    try:
        doc = Document(word_path)
        text_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_lines.append(text)
        
        # Также извлекаем данные из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Пропускаем пустые строки
                    text_lines.append(" | ".join(row_data))
        
        return text_lines
    except Exception as e:
        print(f"Ошибка при чтении Word файла: {e}")
        return []


def parse_priest_data(text_lines: list) -> list:
    """
    Парсинг данных о священниках из текста
    
    Это базовая реализация. В зависимости от формата Word документа
    может потребоваться доработка.
    
    Args:
        text_lines: Список строк текста
    
    Returns:
        Список словарей с данными о священниках
    """
    priests = []
    current_priest = {}
    
    for line in text_lines:
        line = line.strip()
        if not line:
            continue
        
        # Попытка извлечь данные из строки
        # Это примерный парсер, нужно адаптировать под ваш формат
        
        # Если строка содержит разделитель таблицы
        if " | " in line:
            parts = [p.strip() for p in line.split(" | ")]
            if len(parts) >= 2:
                # Предполагаем, что это таблица
                # Первая строка может быть заголовками
                if "имя" in parts[0].lower() or "фамилия" in parts[0].lower():
                    continue  # Пропускаем заголовки
                
                # Пытаемся извлечь данные
                if len(parts) >= 2:
                    name_parts = parts[0].split()
                    if len(name_parts) >= 2:
                        current_priest = {
                            'name': name_parts[0],
                            'surname': ' '.join(name_parts[1:]),
                        }
                        # Пытаемся извлечь остальные данные
                        if len(parts) > 1:
                            current_priest['status'] = parts[1] if len(parts) > 1 else ''
                        if len(parts) > 2:
                            current_priest['service_place'] = parts[2] if len(parts) > 2 else ''
                        
                        priests.append(current_priest.copy())
                        current_priest = {}
        
        # Попытка найти имя и фамилию в обычном тексте
        # Ищем паттерны типа "Иванов Иван" или "Иван Иванов"
        name_pattern = r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ][а-яё]+)'
        match = re.search(name_pattern, line)
        if match:
            # Предполагаем, что первое слово - фамилия, второе - имя
            surname, name = match.groups()
            current_priest = {
                'name': name,
                'surname': surname,
            }
            
            # Пытаемся найти статус
            for status in ['Протоиерей', 'Иерей', 'Диакон', 'Протодиакон']:
                if status in line:
                    current_priest['status'] = status
                    break
            
            # Пытаемся найти дату
            date_pattern = r'(\d{1,2}[./]\d{1,2}[./]\d{4})'
            date_match = re.search(date_pattern, line)
            if date_match:
                current_priest['birth_date'] = date_match.group(1)
            
            priests.append(current_priest.copy())
            current_priest = {}
    
    return priests


def create_excel_from_data(priests: list, output_path: str):
    """
    Создание Excel файла из данных о священниках
    
    Args:
        priests: Список словарей с данными
        output_path: Путь для сохранения Excel файла
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Священники"
    
    # Заголовки
    headers = [
        "Имя", "Фамилия", "Дата рождения", "Место рождения",
        "Статус", "Дата рукоположения", "Место служения",
        "Образование", "Последняя награда"
    ]
    
    # Стили заголовков
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
    
    # Заполнение данных
    for row_num, priest in enumerate(priests, 2):
        ws.cell(row=row_num, column=1, value=priest.get('name', ''))
        ws.cell(row=row_num, column=2, value=priest.get('surname', ''))
        ws.cell(row=row_num, column=3, value=priest.get('birth_date', ''))
        ws.cell(row=row_num, column=4, value=priest.get('birth_place', ''))
        ws.cell(row=row_num, column=5, value=priest.get('status', ''))
        ws.cell(row=row_num, column=6, value=priest.get('ordination_date', ''))
        ws.cell(row=row_num, column=7, value=priest.get('service_place', ''))
        ws.cell(row=row_num, column=8, value=priest.get('education', ''))
        ws.cell(row=row_num, column=9, value=priest.get('last_reward', ''))
    
    # Настройка ширины колонок
    widths = [15, 15, 18, 20, 15, 20, 35, 30, 25]
    for col_num, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + col_num)].width = width
    
    wb.save(output_path)
    print(f"✓ Excel файл создан: {output_path}")


def convert_word_to_excel(word_path: str, excel_path: str = None):
    """
    Главная функция конвертации Word → Excel
    
    Args:
        word_path: Путь к Word файлу
        excel_path: Путь для сохранения Excel файла (опционально)
    """
    if excel_path is None:
        excel_path = word_path.replace('.docx', '.xlsx').replace('.doc', '.xlsx')
    
    print(f"Чтение Word файла: {word_path}")
    text_lines = extract_text_from_word(word_path)
    
    if not text_lines:
        print("⚠️  Не удалось извлечь данные из Word файла")
        return
    
    print(f"Извлечено строк: {len(text_lines)}")
    print("Парсинг данных...")
    
    priests = parse_priest_data(text_lines)
    
    if not priests:
        print("⚠️  Не удалось распарсить данные о священниках")
        print("\nВНИМАНИЕ: Автоматический парсинг может не работать с вашим форматом Word.")
        print("Рекомендуется:")
        print("1. Открыть Word файл в Excel напрямую (если это таблица)")
        print("2. Или вручную скопировать данные в шаблон Excel")
        print("3. Использовать команду: python excel_template.py для создания шаблона")
        return
    
    print(f"Найдено священников: {len(priests)}")
    print("Создание Excel файла...")
    
    create_excel_from_data(priests, excel_path)
    print(f"\n✓ Конвертация завершена!")
    print(f"Файл сохранен: {excel_path}")
    print("\n⚠️  ВНИМАНИЕ: Проверьте данные в Excel файле перед импортом!")
    print("Автоматический парсинг может содержать ошибки.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Использование: python convert_word_to_excel.py <путь_к_word_файлу> [путь_к_excel_файлу]")
        print("\nПример:")
        print("  python convert_word_to_excel.py data.docx")
        print("  python convert_word_to_excel.py data.docx output.xlsx")
        sys.exit(1)
    
    word_file = sys.argv[1]
    excel_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    convert_word_to_excel(word_file, excel_file)
